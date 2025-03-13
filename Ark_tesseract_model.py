import fitz
import cv2
import pytesseract
from pytesseract import Output
import matplotlib.pyplot as plt
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re
import nltk
from nltk.corpus import words
import numpy as np

# Tải bộ dữ liệu từ vựng tiếng Anh
nltk.download('words')
nltk.download('punkt')
nltk.download('punkt_tab')
english_words = set(words.words())

# Cấu hình Tesseract OCR
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Hàm tải mô hình VietOCR
def tai_vietocr():
    config = Cfg.load_config_from_name('vgg_transformer')
    config['device'] = 'cpu'  # Hoặc 'cuda' nếu sử dụng GPU
    config['predictor']['beamsearch'] = True
    return Predictor(config)

vietocr = tai_vietocr()

# Kiểm tra chuỗi có ký tự đặc biệt hay không
def co_ky_tu_dac_biet(tu):
    pattern = re.compile(r'[^\w\s]', re.UNICODE)  # Tìm ký tự không phải là chữ cái, số hoặc khoảng trắng
    return bool(pattern.search(tu))

# Kiểm tra chuỗi có phải là từ tiếng Anh không
def la_tieng_anh(tu):
    return tu.lower() in english_words

# Hiển thị ảnh và các bounding box
def hien_thi_anh_va_box(anh, cac_box):
    # Sắp xếp các bounding box theo vị trí y, x
    cac_box.sort(key=lambda x: (x[1], x[0]))  # Sắp xếp theo y, sau đó theo x
    
    for (x, y, w, h) in cac_box:
        cv2.rectangle(anh, (x, y), (x + w, y + h), (0, 255, 0), 2)
    
    plt.imshow(cv2.cvtColor(anh, cv2.COLOR_BGR2RGB))
    plt.axis("off")
    plt.show()

def tien_xu_li_anh_pdf(anh):
    # Chuyển ảnh thành ảnh xám
    xam = cv2.cvtColor(anh, cv2.COLOR_BGR2GRAY)
    
    # Chuyển ảnh xám thành ảnh nhị phân (ngược màu: văn bản trắng trên nền đen)
    _, anh_nhi_phan = cv2.threshold(xam, 130, 255, cv2.THRESH_BINARY_INV)

    # Làm mờ ảnh để giảm nhiễu (sử dụng GaussianBlur thay vì bilateralFilter)
    blur = cv2.GaussianBlur(anh_nhi_phan, (5, 5), 0)

    # Padding để giữ kích thước ảnh không thay đổi
    padding_size = 1  # Đệm 1 pixel xung quanh
    anh_padding = cv2.copyMakeBorder(blur, padding_size, padding_size, padding_size, padding_size, cv2.BORDER_CONSTANT, value=0)

    # Áp dụng bộ lọc Laplacian
    laplacian = cv2.Laplacian(anh_padding, cv2.CV_64F)

    # Chuyển giá trị Laplacian về dạng 8-bit (ảnh đen trắng)
    anh_lam_net = cv2.convertScaleAbs(laplacian)

    # Loại bỏ padding để giữ nguyên kích thước ban đầu
    anh_lam_net = anh_lam_net[padding_size:-padding_size, padding_size:-padding_size]

    return anh_lam_net

def tien_xu_li_anh_thuong(anh):
    # Chuyển ảnh thành ảnh xám
    xam = cv2.cvtColor(anh, cv2.COLOR_BGR2GRAY)
    
    # Chuyển ảnh xám thành ảnh nhị phân
    _, anh_nhi_phan = cv2.threshold(xam, 130, 255, cv2.THRESH_BINARY_INV)
    
    # Áp dụng bộ lọc Laplacian
    laplacian = cv2.Laplacian(anh_nhi_phan, cv2.CV_64F)
    anh_lam_net = cv2.convertScaleAbs(laplacian)  # Chuyển đổi lại giá trị
    
    # Làm mờ nhẹ để giảm nhiễu
    blur = cv2.bilateralFilter(anh_lam_net, 9, 75, 75)
    
    return blur

# Xử lý ảnh và trích xuất văn bản
def xu_ly_anh(anh, duoi_file):
    print("Đang xử lý hình ảnh...")
    if duoi_file == ".pdf":
        anh_nhi_phan = tien_xu_li_anh_pdf(anh)
    else: anh_nhi_phan= tien_xu_li_anh_thuong(anh)
    ket_qua_tesseract = pytesseract.image_to_data(anh_nhi_phan, output_type=Output.DICT)
    van_ban = pytesseract.image_to_string(anh_nhi_phan, lang='eng+vie')

    cac_vung_box = []
    cac_van_ban = []
    for i in range(len(ket_qua_tesseract['text'])):
        (x, y, w, h) = (ket_qua_tesseract['left'][i], ket_qua_tesseract['top'][i], ket_qua_tesseract['width'][i], ket_qua_tesseract['height'][i])
        tu_tesseract = ket_qua_tesseract['text'][i]

        if tu_tesseract.strip() == "":
            continue

        cac_vung_box.append((x, y, w, h))

        # Xử lý văn bản: Bỏ qua VietOCR nếu có ký tự đặc biệt hoặc là từ tiếng Anh
        if co_ky_tu_dac_biet(tu_tesseract) or la_tieng_anh(tu_tesseract):
            tu_vietocr = tu_tesseract
        else:
            # Xử lý bằng VietOCR nếu không phải ký tự đặc biệt hoặc từ tiếng Anh
            vung_anh = anh[y:y + h, x:x + w]
            anh_roi = Image.fromarray(cv2.cvtColor(vung_anh, cv2.COLOR_BGR2RGB))
            tu_vietocr = vietocr.predict(anh_roi)

        cac_van_ban.append((tu_tesseract, tu_vietocr))

    # Thay thế văn bản Tesseract bằng văn bản VietOCR nếu cần
    for van_tess, van_vietocr in cac_van_ban:
        if van_tess and van_vietocr:
            van_ban = van_ban.replace(van_tess, van_vietocr, 1)

    hien_thi_anh_va_box(anh, cac_vung_box)
    return van_ban

# Xử lý PDF hoặc hình ảnh
def xu_ly_pdf_hoac_anh(duong_dan_vao):
    if not os.path.exists(duong_dan_vao):
        print("Tệp không tồn tại.")
        return

    duoi_file = os.path.splitext(duong_dan_vao)[1].lower()
    cac_van_ban = []

    if duoi_file == '.pdf':
        try:
            tai_lieu = fitz.open(duong_dan_vao)
            cac_trang = []

            for so_trang in range(len(tai_lieu)):
                print(f"Đang xử lý trang {so_trang + 1}...")
                trang = tai_lieu[so_trang]

                # Chuyển đổi trang PDF thành hình ảnh
                pixmap = trang.get_pixmap(dpi=300)
                anh = np.array(Image.frombytes(
                    "RGB", [pixmap.width, pixmap.height], pixmap.samples
                ))
                cac_trang.append(anh)

            # Xử lý từng trang
            for so_trang, anh_trang in enumerate(cac_trang):
                van_ban_trang = xu_ly_anh(anh_trang,duoi_file)
                print(f"Văn bản từ trang {so_trang + 1}:{van_ban_trang}\n")
                cac_van_ban.append(van_ban_trang)
        except Exception as e:
            print(f"Lỗi khi xử lý PDF: {e}")
            return

    else:
        try:
            anh = cv2.imread(duong_dan_vao)
            if anh is None:
                print("Không thể đọc ảnh.")
                return

            # Xử lý ảnh
            van_ban_anh = xu_ly_anh(anh,duoi_file)
            print(f"Văn bản từ ảnh:\n{van_ban_anh}\n")
            cac_van_ban.append(van_ban_anh)
        except Exception as e:
            print(f"Lỗi khi xử lý ảnh: {e}")
            return

    # Ghép các đoạn văn và lưu vào file Word
    van_ban_hoan_thanh = "\n".join(cac_van_ban)
    hoan_thanh_van_ban(van_ban_hoan_thanh, duong_dan_vao)

# Lưu văn bản vào Word
def hoan_thanh_van_ban(van_ban_hoan_chinh, duong_dan_anh):
    print("Đang lưu văn bản vào file Word...")

    # Lấy tên file ảnh và tạo đường dẫn lưu file Word với tên ảnh
    ten_anh = os.path.splitext(os.path.basename(duong_dan_anh))[0]
    
    # Cập nhật đường dẫn lưu file Word vào thư mục cần lưu
    duong_dan_luu = os.path.join(r"F:\Docx", f"{ten_anh}.docx") # Ví dụ ảnh tên 1 vie.png F:\Docx\1vie.docx
    
    # Tạo tài liệu mới
    doc = Document()

    # Thiết lập kiểu font mặc định cho tài liệu
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # Thêm toàn bộ văn bản vào tài liệu Word
    doan_van = doc.add_paragraph(van_ban_hoan_chinh.strip())
    doan_van.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Thiết lập font cho từng Run (đề phòng trường hợp một đoạn có nhiều Run)
    for run in doan_van.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Lưu văn bản vào file với tên theo ảnh trong thư mục mới
    doc.save(duong_dan_luu)
    print(f"Lưu thành công: {duong_dan_luu}")

    # Xóa khoảng trắng không cần thiết trong các đoạn văn (nếu cần)
    xoa_khoang_trang(duong_dan_luu)

# Hàm xóa khoảng trắng không cần thiết trong các đoạn văn
def xoa_khoang_trang(duong_dan_file):
    print(f"Đang xóa khoảng trắng không cần thiết trong file {duong_dan_file}...")
    
    # Mở lại file Word đã lưu
    doc = Document(duong_dan_file)
    
    # Lấy danh sách các đoạn văn trong tài liệu
    paragraphs = doc.paragraphs
    
    # Duyệt qua các đoạn văn và xóa khoảng trắng không cần thiết
    for para in paragraphs:
        # Chỉ xử lý các đoạn không trống (có nội dung)
        if para.text.strip():  # Kiểm tra nếu đoạn văn không trống
            # Xóa khoảng trắng ở đầu và cuối đoạn
            para.text = para.text.strip()
    
    # Lưu lại file Word sau khi đã xử lý
    doc.save(duong_dan_file)
    print(f"Đã xóa khoảng trắng và lưu lại file: {duong_dan_file}")

# Đường dẫn tệp PDF hoặc ảnh đầu vào
duong_dan = r"F:\OCR\Vie image data\1 vie.jpg"
xu_ly_pdf_hoac_anh(duong_dan)